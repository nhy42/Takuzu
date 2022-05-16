from os import listdir
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE



def processParam(pString):
    pString = pString.split("(")[1]
    pString = ")".join(pString.split(")")[:-1])
    if pString == "":
        return []
    elif "," not in pString:
        return [pString]
    else:
        listToRet = pString.split(",")
        for i in range(len(listToRet)):
            listToRet[i] = listToRet[i].strip()
        return listToRet


pyFileList = []
for e in listdir("files\\"):
    if e.split(".")[-1] == "h":
        pyFileList.append(e)
print(pyFileList)

document = Document()
document.add_heading("Titre", 0)
for filename in pyFileList:
    document.add_heading(filename, level=1)

    f = open("files\\" + filename, "r")
    lines = f.readlines()
    for line in lines:
        line = line.split(" ")
        #if line[0] == "def":
        if line[0].strip() != "":
            document.add_heading(" ".join(line[:]), level=2)
            document.add_paragraph('Desc', style="Caption")
            document.add_paragraph('')
        #    paramList = processParam(" ".join(line))
        #    for param in paramList:
        #        p.add_run(param + " <> : \n")
        #    p.add_run("return <> : \n")
    f.close()

document.save("DOCUMENTATION.docx")
